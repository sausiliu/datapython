import time
import xlwt
import xlrd
import pymysql
import pandas as pd
from docx import Document


def create_paragraph(sqls, document):
    # print(sql)

    # Add paragraph
    text = df.at[1, 'set_time'] # 获取特定的值
    document.add_paragraph(u'在这里可以添加文本:' + str(text) +
                           u' :添加文本结束\n')

    for sql in sqls:
        db = pymysql.connect(
            host='10.10.30.25',
            user='root',
            passwd='1',
            port=3306,
            charset='utf8')

        db.close()
    print(text)


def create_table(sqls, document):
    # print(sql)

    # 创建word table
    # 读取mysql数据，并且组合放入pandas中
    first = True
    for sql in sqls:
        db = pymysql.connect(
            host='10.10.30.25',
            user='root',
            passwd='1',
            port=3306,
            charset='utf8')

        if first:
            df = pd.read_sql(sql, con=db)
        else:
            df = pd.concat([df, pd.read_sql(sql, con=db)], sort=False)

        db.close()
        first = False

    print(df)

    table = document.add_table(df.shape[0] + 1, df.shape[1], style='Table Grid')

    # Add table
    # add the header rows.
    for j in range(df.shape[-1]):
        table.cell(0, j).text = df.columns[j]

    # add the rest of the data frame
    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            table.cell(i + 1, j).text = str(df.values[i, j])

    # 行求和df.iloc[0,0:].sum()


def paste_table_word(path):

    readbook = xlrd.open_workbook(path)
    sheet = readbook.sheet_by_index(0)
    #sheet = readbook.sheet_by_name()

    ncols = sheet.ncols
    nrows = sheet.nrows
    print("create a table {} x {} ".format(ncols, nrows))

    document = Document()
    # document.add_heading(u'我的一个新文档', 0)

    # export_data_word(sqls)
    document.add_paragraph("This style is : ")
    table = document.add_table(rows=nrows, cols=ncols, style='Table Grid')

    for r in range(nrows):
        for c in range(ncols):
            #print(sheet.cell(r, c))
            cell = table.cell(r, c)
            cell.text = u'' + str(sheet.cell(r, c).value)

    #table.add_row()
    table.add_column(10)

    document.save('./test.docx')


def export_data_excel(sqls):

    # time.sleep(1)
    # Create xlsx
    excel = xlwt.Workbook()
    sheet = excel.add_sheet(u'sheet1', cell_overwrite_ok=True)

    global_c = 0

    # 执行多个sql查询语句，并写入excel
    for sql in sqls:
        db = pymysql.connect(
            host='10.10.30.25',
            user='root',
            passwd='1',
            port=3306,
            charset='utf8')

        cursor = db.cursor()  # 建立游标
        cursor.execute(sql)
        data = cursor.fetchall()
        fields = cursor.description
        # print("len fields: {}".format(len(fields)))
        # print("len data: {}".format(len(data)))

        db.close()
        # 写入行名称
        for i in range(len(fields)):
            # print(fields[i][0])
            sheet.write(0, i+global_c, fields[i][0])

        # 写内容
        for r in range(len(data)):
            for c in range(len(fields)):
                #  print(sensor_data[r][c])
                sheet.write(1+r, c+global_c, data[r][c])

        global_c = global_c + len(fields)
    # sensor_data = cursor.fetchone()

    excel.save('test.xls')


def set_style(name, height, bold=False):

    style = xlwt.XFStyle()  # 初始化样式
    font = xlwt.Font()
    font.name = name  # 'Times New Roman'
    font.bold = bold
    font.colour_index = 4
    font.height = height
    style.font = font

    return style


if __name__ == "__main__":

    # sheet = excel.add_sheet(u'sheet1', cell_overwrite_ok=True)
    sqls = [
        """SELECT * FROM `sys`.`sys_config` LIMIT 0, 1000"""
    ]

    # export_data_excel(sqls)

    # paste_table_word('./test.xls')

    # Create word
    document = Document()
    create_table(sqls, document)
    document.save('./test.docx')
