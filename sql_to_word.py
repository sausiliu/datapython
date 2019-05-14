import xlwt
import xlrd
import pymysql
import pandas as pd
from decimal import Decimal
from docx import Document
from docx.oxml.ns import qn
# import data_connection_ffm as con

global database


def connect_mysql():
    db = pymysql.connect(
         host='10.211.55.5',
         user='root',
         passwd='1',
         port=3306,
         charset='utf8')
    return db


def sql_to_pandas(sqls):
    first = True
    for sql in sqls:
        if first:
            df = pd.read_sql(sql, con=database)
        else:
            df = pd.concat([df, pd.read_sql(sql, con=database)], sort=False)
        first = False
    df = df.reset_index()
    print(df)
    return df


def create_paragraph1(sqls, document):
    # print(sql)
    df = sql_to_pandas(sqls)
    # Add paragraph
    # text = df.at[0, 'count_interest_all'] # 获取特定的值
    text = df.at[0, 'TOTAL_CONNECTIONS']  # 获取特定的值
    document.add_paragraph(u'在这里可以添加文本:' + str( Decimal((text+text)/2.1).quantize(Decimal("0.00")) ) +
                           u' :添加文本结束\n')
    print(text)


def create_paragraph2(sqls, document):
    df = sql_to_pandas(sqls)

    # Add paragraph
    # text = df.at[0, 'count_interest_all'] # 获取特定的值
    text = df.at[0, 'TOTAL_CONNECTIONS']  # 获取特定的值
    document.add_paragraph(u'22在这里可以添加文本:' + str(text) +
                           u' :添加文本结束\n')
    print(text)


def word_add_table(sqls, document):
    # print(sql)

    # 创建word table
    # 读取mysql数据，并且组合放入pandas中
    df = sql_to_pandas(sqls)

    table = document.add_table(df.shape[0] + 1, df.shape[1], style='Table Grid')

    # Add table
    # add the header rows.
    for j in range(df.shape[-1]):
        table.cell(0, j).text = df.columns[j]

    # add the rest of the data frame
    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            table.cell(i + 1, j).text = str(df.values[i, j])

    # 行求和df.iloc[0,0:].sum()


def paste_table_word(path):
    readbook = xlrd.open_workbook(path)
    sheet = readbook.sheet_by_index(0)
    # sheet = readbook.sheet_by_name()

    ncols = sheet.ncols
    nrows = sheet.nrows
    print("create a table {} x {} ".format(ncols, nrows))

    document = Document()
    # document.add_heading(u'我的一个新文档', 0)

    # export_data_word(sqls)
    document.add_paragraph("This style is : ")
    table = document.add_table(rows=nrows, cols=ncols, style='Table Grid')

    for r in range(nrows):
        for c in range(ncols):
            # print(sheet.cell(r, c))
            cell = table.cell(r, c)
            cell.text = u'' + str(sheet.cell(r, c).value)

    # table.add_row()
    table.add_column(10)

    document.save('./test.docx')


def export_data_excel(sqls, sheet):

    global_c = 0

    cursor = database.cursor()  # 建立游标
    # 执行多个sql查询语句，并写入excel
    for sql in sqls:
        cursor.execute(sql)
        data = cursor.fetchall()
        fields = cursor.description
        # print("len fields: {}".format(len(fields)))
        # print("len data: {}".format(len(data)))

        # 写入行名称
        for i in range(len(fields)):
            # print(fields[i][0])
            sheet.write(0, i+global_c, fields[i][0])

        # 写内容
        for r in range(len(data)):
            for c in range(len(fields)):
                #  print(sensor_data[r][c])
                sheet.write(1+r, c+global_c, data[r][c])

        global_c = global_c + len(fields)


def set_style(name, height, bold=False):
    style = xlwt.XFStyle()  # 初始化样式
    font = xlwt.Font()
    font.name = name  # 'Times New Roman'
    font.bold = bold
    font.colour_index = 4
    font.height = height
    style.font = font

    return style


if __name__ == "__main__":

    # 1. 连接数据库
    database = connect_mysql()

    # 准备sqls

    sqls = [
        """SELECT * FROM `performance_schema`.`hosts` LIMIT 0, 1000""",
        """SELECT * FROM `performance_schema`.`hosts` LIMIT 0, 1000""",
        """SELECT * FROM `performance_schema`.`hosts` LIMIT 0, 1000"""
    ]

    # $$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$
    #      This is a excel demo     $
    # $$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$

    # Create a excel file
    excel = xlwt.Workbook()
    sheet = excel.add_sheet(u'sheet1', cell_overwrite_ok=True)
    export_data_excel(sqls, sheet)
    sheet = excel.add_sheet(u'sheet2', cell_overwrite_ok=True)
    export_data_excel(sqls, sheet)
    excel.save('test.xls')

    # $$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$
    #      This is a word demo      $
    # $$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$
    # Create word
    document = Document()
    # 设置为宋体
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    # 添加文字
    sql1 = [ """SELECT * FROM `performance_schema`.`hosts` LIMIT 0, 1000"""]
    sql2 = [ """SELECT * FROM `performance_schema`.`hosts` LIMIT 0, 1000"""]
    create_paragraph1(sql1, document)
    create_paragraph2(sql2, document)

    # 添加表格
    word_add_table(sqls, document)

    document.save('test.docx')

    # end: 关闭数据库
    database.close()
